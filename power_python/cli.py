# PowerPython
# Copyright (c) 2026 PowerPython contributors
# This file is derived from MATPOWER, Copyright (c) 1996-2025,
# Power Systems Engineering Research Center (PSERC) by Ray Zimmerman, PSERC Cornell.
# Licensed under the 3-clause BSD License (see LICENSE file for details).
# See https://github.com/salorajan/matpower_data_migration for more info.

import sys
import os
import numpy as np
import pandas as pd
import docx
from docx.shared import Pt

# Expose package path
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

from power_python.core.case import PowerCase
from power_python.core.constants import *
from power_python.solvers.runpf import run_power_flow
from power_python.solvers.complex_nr import run_complex_nr
from power_python.solvers.complex_nr_3p import run_complex_nr_3p
from power_python.solvers.rundcpf import run_dc_pf
from power_python.solvers.dcopf import run_dc_opf
from power_python.solvers.acopf import run_ac_opf
from power_python.solvers.sdp_opf import run_sdp_opf
from power_python.solvers.uopf import run_uopf
from power_python.solvers.cpf import run_cpf
from power_python.solvers.hepf import run_hepf
from power_python.solvers.se import run_state_estimation
from power_python.solvers.radial_pf import run_radial_pf
from power_python.solvers.pf_3p import run_3p_pf
from power_python.solvers.opf_3p import run_3p_opf
from power_python.solvers.mp_opf import run_mp_opf
from power_python.solvers.stochastic_opf import run_stochastic_opf
from power_python.solvers.market import run_market_auction
from power_python.solvers.var_planning import run_var_planning
from power_python.solvers.contingency import run_contingency_analysis
from power_python.solvers.sc_opf import run_sc_opf

from power_python.network.sensitivity import make_ptdf, make_lodf
from power_python.utils.audit import calculate_system_balance
from power_python.utils.lmp_decomp import decompose_dc_lmp
from power_python.utils.costs import calculate_total_cost

# Root data directory mapping
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
CASE_DIR = os.path.join(BASE_DIR, "outputs", "json")

HELP_DICT = {
    "acpf": """AC Power Flow (Newton-Raphson)
Usage: acpf <case_id> [accuracy] [excel|csv|docx]
Example: acpf case14 1e-5 excel
Physics/Algorithm: 
Solves AC power flow using the standard Newton-Raphson iteration in polar coordinates. 
Constructs a coupled Jacobian matrix of active/reactive mismatches with respect to voltage angle and magnitude, updating state variables quadratically in 3-5 iterations.""",
    
    "gausspf": """AC Power Flow (Gauss-Seidel)
Usage: gausspf <case_id> [accuracy] [excel|csv|docx]
Example: gausspf case9 1e-4 docx
Physics/Algorithm:
Solves AC power flow using legacy Gauss-Seidel nodal propagation.
It solves voltage variables sequentially (Seidel mode) based on the latest iteration estimates. Slow convergence, but requires no Jacobian construction.""",
    
    "fdpf": """AC Power Flow (Fast Decoupled)
Usage: fdpf <case_id> [accuracy] [excel|csv|docx]
Example: fdpf case30 1e-5 csv
Physics/Algorithm:
Solves AC power flow using the Fast Decoupled method (XB formulation).
Decouples voltage magnitudes and angles by neglecting conductance matrices, using constant B' and B'' susceptance matrices. Highly efficient per-iteration scaling.""",
    
    "dcpf": """DC Power Flow (Linearized)
Usage: dcpf <case_id> [accuracy] [excel|csv|docx]
Example: dcpf case118 excel
Physics/Algorithm:
Solves linearized DC power flow by assuming voltage magnitudes are flat at 1.0 p.u., neglecting resistance (R=0) and shunt admittances, resulting in a single linear system solve.""",
    
    "hepf": """Holomorphic Embedding Power Flow (HEPF)
Usage: hepf <case_id> [accuracy] [excel|csv|docx]
Example: hepf case9 1e-6 docx
Physics/Algorithm:
A non-iterative power flow solver based on analytic continuation. 
Formulates a holomorphic parameterization of the network and evaluates physical branch voltage functions via Padé Approximants, guaranteeing discovery of the high-voltage solution without starting guesses.""",
    
    "cnr": """Complex Newton-Raphson (CVNR)
Usage: cnr <case_id> [accuracy] [excel|csv|docx]
Example: cnr case14 1e-8 excel
Physics/Algorithm:
Solves power flow directly in the complex variable plane using Wirtinger calculus derivatives. 
Bypasses polar decomposition, solving a complex mismatch system of size 2N x 2N.""",
    
    "pf3p": """Three-Phase Unbalanced Power Flow
Usage: pf3p <case_id> [accuracy] [excel|csv|docx]
Example: pf3p case3p_a 1e-5 excel
Physics/Algorithm:
Solves unbalanced 3-phase power flow using the Z-bus iterative method.
Accommodates coupling matrices across phases a-b-c, suitable for distribution system modeling.""",
    
    "radial": """Radial Backward-Forward Sweep Power Flow (BFS)
Usage: radial <case_id> [accuracy] [excel|csv|docx]
Example: radial case33bw 1e-6 csv
Physics/Algorithm:
Backward-Forward Sweep (BFS) solver optimized for radial distribution feeders.
Calculates branch currents backward from leaves to root, then updates bus voltages forward from root to leaves.""",
    
    "dcopf": """DC Optimal Power Flow (DC-OPF)
Usage: dcopf <case_id> [accuracy] [excel|csv|docx]
Example: dcopf case9 excel
Physics/Algorithm:
Optimizes generator active power outputs to minimize costs under transmission line limit constraints. 
Uses linearized DC power flow equations co-optimizing spinning reserves if defined.""",
    
    "acopf": """AC Optimal Power Flow (AC-OPF)
Usage: acopf <case_id> [accuracy] [excel|csv|docx]
Example: acopf case14 docx
Physics/Algorithm:
Solves the non-linear, non-convex AC Optimal Power Flow utilizing SciPy SLSQP.
Minimizes generation cost while co-optimizing active/reactive dispatch, keeping voltage profiles within bounds.""",
    
    "sdpopf": """Convex SDP Optimal Power Flow
Usage: sdpopf <case_id> [accuracy] [excel|csv|docx]
Example: sdpopf case9 excel
Physics/Algorithm:
Formulates the non-convex AC-OPF as a convex Semidefinite Programming (SDP) relaxation (Lavaei & Low method).
Relaxes the rank-1 constraint on the voltage matrix. Evaluates global optimality bounds directly via eigenvalue gap.""",
    
    "uopf": """Unit Decommitment and OPF (UOPF)
Usage: uopf <case_id> [accuracy] [excel|csv|docx]
Example: uopf case9 1e-5 excel
Physics/Algorithm:
Determines optimal unit decommitment status and co-optimized generator dispatch using heuristic search.
Progressively shuts down uneconomical generators to minimize total system cost while satisfying grid constraints.""",
    
    "scopf": """Security-Constrained OPF (SC-OPF)
Usage: scopf <case_id> [accuracy] [excel|csv|docx]
Example: scopf case30 excel
Physics/Algorithm:
Solves security-constrained active power dispatch.
Guarantees grid operation remains secure and line flows do not exceed emergency limits under any single N-1 transmission outage.""",
    
    "opf3p": """Three-Phase Unbalanced AC-OPF
Usage: opf3p <case_id> [accuracy] [excel|csv|docx]
Example: opf3p case3p_a excel
Physics/Algorithm:
Solves unbalanced three-phase AC Optimal Power Flow.
Optimizes generator dispatch at the phase level (a-b-c) to minimize unbalanced distribution costs under operational constraints.""",
    
    "mpopf": """Multi-Period OPF with Energy Storage (MPOPF)
Usage: mpopf <case_id> [accuracy] [excel|csv|docx]
Example: mpopf case9 excel
Physics/Algorithm:
Co-optimizes generator active power dispatch and battery charging/discharging profiles over a multi-period horizon.
Integrates battery state-of-charge constraints and generator ramping limits.""",
    
    "stopf": """Stochastic DC Optimal Power Flow
Usage: stopf <case_id> [accuracy] [excel|csv|docx]
Example: stopf case9 csv
Physics/Algorithm:
Solves stochastic active dispatch optimization.
Minimizes the expected generation costs over multiple probabilistic renewable generation scenarios (wind/solar).""",
    
    "market": """Smart Market Auction Simulation
Usage: market <case_id> [accuracy] [excel|csv|docx]
Example: market case9 excel
Physics/Algorithm:
Simulates a competitive electricity market auction.
Clears generator step offers and load bids under network flow constraints, calculating cleared power and turnover.""",
    
    "varplan": """VAr Planning / Optimal Capacitor Placement
Usage: varplan <case_id> [accuracy] [excel|csv|docx]
Example: varplan case14 docx
Physics/Algorithm:
Solves optimal reactive power planning using AC-OPF models.
Places shunt susceptance (capacitors) at optimal candidate buses to keep voltages within limits at minimum installation cost.""",
    
    "contingency": """N-1 Contingency Analysis Screening
Usage: contingency <case_id> [accuracy] [excel|csv|docx]
Example: contingency case14 csv
Physics/Algorithm:
Performs systematic screening of single-line outage contingencies using PTDF and LODF sensitivity factors.
Quickly identifies potential line overload violations across all single line outages.""",
    
    "se": """Weighted Least Squares State Estimation (SE)
Usage: se <case_id> [accuracy] [excel|csv|docx]
Example: se case9 1e-6 docx
Physics/Algorithm:
Solves power system state estimation using Weighted Least Squares (WLS).
Filters noise from redundant measurements (power, voltage, flow) to estimate the true state variables (voltages).""",
    
    "cpf": """Continuation Power Flow (CPF)
Usage: cpf <case_id> [accuracy] [excel|csv|docx]
Example: cpf case9 excel
Physics/Algorithm:
Traces the grid's voltage loadability profile (PV curve) using predictor-corrector continuation steps.
Maintains numeric stability near bifurcation, identifying the maximum loadability nose point.""",
    
    "audit": """Physical Power Balance Audit
Usage: audit <case_id> [accuracy] [excel|csv|docx]
Example: audit case9 docx
Physics/Algorithm:
Checks physical power conservation in the network.
Computes absolute active and reactive residuals (Total Generation - Total Load - Losses) to verify case validity.""",
    
    "lmp": """LMP Energy/Congestion Decomposition
Usage: lmp <case_id> [accuracy] [excel|csv|docx]
Example: lmp case9 excel
Physics/Algorithm:
Decomposes Locational Marginal Prices (LMP) from DC-OPF dual variables.
Partitions nodal marginal costs into Energy and Congestion components to highlight transmission congestion pricing."""
}

def load_case_by_id(case_id):
    full_name = case_id if case_id.startswith("case") else f"case{case_id}"
    # Remove file extension if user typed it
    if full_name.endswith(".json"):
        full_name = full_name[:-5]
    elif full_name.endswith(".xlsx"):
        full_name = full_name[:-5]
        
    case_path = os.path.join(CASE_DIR, f"{full_name}.json")
    if not os.path.exists(case_path):
        print(f"Error: Case file not found at {case_path}")
        return None
        
    case = PowerCase()
    case.load_from_json(case_path)
    return case

def export_results_excel(case, filename, analysis, extra_results=None):
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill
    
    # Re-calculate generator PG and QG outputs first to ensure they are up to date in the exported case
    case.update_generator_power()
    
    with pd.ExcelWriter(filename, engine='openpyxl') as writer:
        # Write General sheet if it has baseMVA
        if hasattr(case, 'baseMVA'):
            gen_df = pd.DataFrame({"baseMVA": [case.baseMVA]})
            gen_df.to_excel(writer, sheet_name='General', index=False)
            
        if hasattr(case, 'bus') and case.bus is not None and len(case.bus) > 0:
            # Calculate generation per bus
            gen_p = np.zeros(len(case.bus))
            gen_q = np.zeros(len(case.bus))
            for i in range(len(case.gen)):
                if case.gen[i, GEN_STATUS] > 0:
                    bus_idx = int(case.gen[i, GEN_BUS])
                    gen_p[bus_idx] += case.gen[i, PG]
                    gen_q[bus_idx] += case.gen[i, QG]
            
            # Bus sheet in MATLAB MATPOWER format, enhanced with active/reactive generation
            bus_data = {
                "BUS_I": case.external_bus_ids.astype(int),
                "TYPE": case.bus[:, BUS_TYPE].astype(int),
                "VM": case.bus[:, VM],
                "VA": case.bus[:, VA],
                "PD": case.bus[:, PD],
                "QD": case.bus[:, QD],
                "PG": gen_p,
                "QG": gen_q,
                "GS": case.bus[:, GS],
                "BS": case.bus[:, BS],
                "BUS_AREA": case.bus[:, BUS_AREA].astype(int),
                "BASE_KV": case.bus[:, BASE_KV],
                "ZONE": case.bus[:, ZONE].astype(int),
                "VMAX": case.bus[:, VMAX],
                "VMIN": case.bus[:, VMIN],
            }
            if case.bus.shape[1] > 13:
                # Add OPF variables if they exist in the matrix
                bus_data["LAM_P"] = case.bus[:, LAM_P]
                bus_data["LAM_Q"] = case.bus[:, LAM_Q]
                bus_data["MU_VMAX"] = case.bus[:, MU_VMAX]
                bus_data["MU_VMIN"] = case.bus[:, MU_VMIN]
                
            bus_df = pd.DataFrame(bus_data)
            bus_df.to_excel(writer, sheet_name='Bus', index=False)
            
            # Generator sheet in MATLAB MATPOWER format
            gen_bus_ids = [int(case.external_bus_ids[int(b)]) for b in case.gen[:, GEN_BUS]]
            gen_data = {
                "GEN_BUS": gen_bus_ids,
                "PG": case.gen[:, PG],
                "QG": case.gen[:, QG],
                "QMAX": case.gen[:, QMAX],
                "QMIN": case.gen[:, QMIN],
                "VG": case.gen[:, VG],
                "MBASE": case.gen[:, MBASE],
                "GEN_STATUS": case.gen[:, GEN_STATUS].astype(int),
                "PMAX": case.gen[:, PMAX],
                "PMIN": case.gen[:, PMIN],
                "PC1": case.gen[:, PC1],
                "PC2": case.gen[:, PC2],
                "QC1MIN": case.gen[:, QC1MIN],
                "QC1MAX": case.gen[:, QC1MAX],
                "QC2MIN": case.gen[:, QC2MIN],
                "QC2MAX": case.gen[:, QC2MAX],
                "RAMP_AGC": case.gen[:, RAMP_AGC],
                "RAMP_10": case.gen[:, RAMP_10],
                "RAMP_30": case.gen[:, RAMP_30],
                "RAMP_Q": case.gen[:, RAMP_Q],
                "APF": case.gen[:, APF],
            }
            if case.gen.shape[1] > 21:
                gen_data["MU_PMAX"] = case.gen[:, MU_PMAX]
                gen_data["MU_PMIN"] = case.gen[:, MU_PMIN]
                gen_data["MU_QMAX"] = case.gen[:, MU_QMAX]
                gen_data["MU_QMIN"] = case.gen[:, MU_QMIN]
                
            gen_df = pd.DataFrame(gen_data)
            gen_df.to_excel(writer, sheet_name='Generator', index=False)
            
            # Branch sheet in MATLAB MATPOWER format
            f_bus_ids = [int(case.external_bus_ids[int(b)]) for b in case.branch[:, F_BUS]]
            t_bus_ids = [int(case.external_bus_ids[int(b)]) for b in case.branch[:, T_BUS]]
            branch_data = {
                "F_BUS": f_bus_ids,
                "T_BUS": t_bus_ids,
                "BR_R": case.branch[:, BR_R],
                "BR_X": case.branch[:, BR_X],
                "BR_B": case.branch[:, BR_B],
                "RATE_A": case.branch[:, RATE_A],
                "RATE_B": case.branch[:, RATE_B],
                "RATE_C": case.branch[:, RATE_C],
                "TAP": case.branch[:, TAP],
                "SHIFT": case.branch[:, SHIFT],
                "BR_STATUS": case.branch[:, BR_STATUS].astype(int),
                "ANGMIN": case.branch[:, ANGMIN],
                "ANGMAX": case.branch[:, ANGMAX],
                "PF": case.branch[:, PF],
                "QF": case.branch[:, QF],
                "PT": case.branch[:, PT],
                "QT": case.branch[:, QT],
            }
            if case.branch.shape[1] > 17:
                branch_data["MU_SF"] = case.branch[:, MU_SF]
                branch_data["MU_ST"] = case.branch[:, MU_ST]
                branch_data["MU_ANGMIN"] = case.branch[:, MU_ANGMIN]
                branch_data["MU_ANGMAX"] = case.branch[:, MU_ANGMAX]
                
            branch_df = pd.DataFrame(branch_data)
            branch_df.to_excel(writer, sheet_name='Branch', index=False)
            
            # Line Flows sheet with detailed flow and loss data (to/fro flows, individual line losses, and total losses)
            p_from = case.branch[:, PF]
            q_from = case.branch[:, QF]
            p_to = case.branch[:, PT]
            q_to = case.branch[:, QT]
            p_loss = p_from + p_to
            q_loss = q_from + q_to
            
            line_flows_data = {
                "F_BUS": f_bus_ids,
                "T_BUS": t_bus_ids,
                "P_FROM_TO": p_from,
                "Q_FROM_TO": q_from,
                "P_TO_FROM": p_to,
                "Q_TO_FROM": q_to,
                "P_LOSS": p_loss,
                "Q_LOSS": q_loss
            }
            line_flows_df = pd.DataFrame(line_flows_data)
            
            # Append Total Row
            total_row = {
                "F_BUS": "TOTAL",
                "T_BUS": "",
                "P_FROM_TO": np.nan,
                "Q_FROM_TO": np.nan,
                "P_TO_FROM": np.nan,
                "Q_TO_FROM": np.nan,
                "P_LOSS": np.sum(p_loss),
                "Q_LOSS": np.sum(q_loss)
            }
            line_flows_df = pd.concat([line_flows_df, pd.DataFrame([total_row])], ignore_index=True)
            line_flows_df.to_excel(writer, sheet_name='Line Flows', index=False)
            
            # Generator Cost sheet if present and not empty
            if hasattr(case, 'gencost') and case.gencost is not None and len(case.gencost) > 0:
                gencost_cols = ["MODEL", "STARTUP", "SHUTDOWN", "NCOST"] + [f"COST_{i}" for i in range(case.gencost.shape[1] - 4)]
                num_cols = min(len(gencost_cols), case.gencost.shape[1])
                gencost_df = pd.DataFrame(case.gencost[:, :num_cols], columns=gencost_cols[:num_cols])
                # Format integer columns
                for col in ["MODEL", "NCOST"]:
                    if col in gencost_df:
                        gencost_df[col] = gencost_df[col].astype(int)
                gencost_df.to_excel(writer, sheet_name='Generator Cost', index=False)
                
        # Write 3-Phase sheets if present
        if hasattr(case, 'bus3p') and case.bus3p is not None and len(case.bus3p) > 0:
            bus3p_df = pd.DataFrame(case.bus3p, columns=["Bus_ID", "Type", "BaseKV", "Vm_a", "Vm_b", "Vm_c", "Va_a", "Va_b", "Va_c"])
            bus3p_df.to_excel(writer, sheet_name='Bus3P', index=False)
            
            if hasattr(case, 'line3p') and case.line3p is not None and len(case.line3p) > 0:
                pd.DataFrame(case.line3p).to_excel(writer, sheet_name='Line3P', index=False)
            if hasattr(case, 'xfmr3p') and case.xfmr3p is not None and len(case.xfmr3p) > 0:
                pd.DataFrame(case.xfmr3p).to_excel(writer, sheet_name='Xfmr3P', index=False)
            if hasattr(case, 'load3p') and case.load3p is not None and len(case.load3p) > 0:
                pd.DataFrame(case.load3p).to_excel(writer, sheet_name='Load3P', index=False)
            if hasattr(case, 'gen3p') and case.gen3p is not None and len(case.gen3p) > 0:
                pd.DataFrame(case.gen3p).to_excel(writer, sheet_name='Gen3P', index=False)
            if hasattr(case, 'lc') and case.lc is not None and len(case.lc) > 0:
                pd.DataFrame(case.lc).to_excel(writer, sheet_name='LineConst', index=False)
                
        # Write Extra / Analysis results
        if extra_results is not None:
            if isinstance(extra_results, pd.DataFrame):
                extra_results.to_excel(writer, sheet_name='Analysis_Results', index=False)
            elif isinstance(extra_results, dict):
                for k, v in extra_results.items():
                    sheet_name = f"Analysis_{k}"[:31]
                    if isinstance(v, pd.DataFrame):
                        v.to_excel(writer, sheet_name=sheet_name, index=False)
                    else:
                        pd.DataFrame(v).to_excel(writer, sheet_name=sheet_name, index=False)

    # Re-open the excel file with openpyxl to apply gorgeous styling & formatting
    wb = openpyxl.load_workbook(filename)
    
    # Styles definition
    hdr_font = Font(name="Segoe UI", size=11, bold=True, color="FFFFFF")
    hdr_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
    data_font = Font(name="Segoe UI", size=10)
    
    center_align = Alignment(horizontal="center", vertical="center")
    right_align = Alignment(horizontal="right", vertical="center")
    left_align = Alignment(horizontal="left", vertical="center")
    
    # Columns that should be centered (integer IDs, types, codes)
    center_cols = {
        "BUS_I", "TYPE", "BUS_AREA", "ZONE", "GEN_BUS", "GEN_STATUS",
        "F_BUS", "T_BUS", "BR_STATUS", "MODEL", "NCOST", "Bus_ID", "Type"
    }
    
    # Columns that should be formatted as 4-decimal floats (voltage magnitudes)
    v_cols = {"VM", "VMAX", "VMIN", "VG", "Vm_a", "Vm_b", "Vm_c"}
    
    # Columns that should be formatted as 2-decimal floats (angles)
    ang_cols = {"VA", "Va_a", "Va_b", "Va_c", "SHIFT", "ANGMIN", "ANGMAX"}
    
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        ws.sheet_view.showGridLines = True
        
        # Determine number of rows and columns
        max_row = ws.max_row
        max_col = ws.max_column
        
        if max_row == 0 or max_col == 0:
            continue
            
        # Freeze headers
        ws.freeze_panes = "A2"
        
        # Read header names
        headers = [ws.cell(row=1, column=c).value for c in range(1, max_col + 1)]
        
        # Style Header Row
        ws.row_dimensions[1].height = 26
        for c in range(1, max_col + 1):
            cell = ws.cell(row=1, column=c)
            cell.font = hdr_font
            cell.fill = hdr_fill
            cell.alignment = center_align
            
        # Style Data Rows
        for r in range(2, max_row + 1):
            ws.row_dimensions[r].height = 20
            for c in range(1, max_col + 1):
                cell = ws.cell(row=r, column=c)
                cell.font = data_font
                
                col_name = headers[c - 1]
                val = cell.value
                
                # Check alignment and number formats
                if col_name in center_cols:
                    cell.alignment = center_align
                    if val is not None:
                        try:
                            cell.value = int(float(val))
                            cell.number_format = "0"
                        except ValueError:
                            pass
                elif isinstance(val, (int, float)):
                    cell.alignment = right_align
                    if col_name in v_cols:
                        cell.number_format = "0.0000"
                    elif col_name in ang_cols:
                        cell.number_format = "0.00"
                    else:
                        cell.number_format = "#,##0.00"
                else:
                    cell.alignment = left_align
                    
        # Special styling for total rows if sheet is 'Line Flows'
        if sheet_name == 'Line Flows':
            total_font = Font(name="Segoe UI", size=10, bold=True)
            for c in range(1, max_col + 1):
                cell = ws.cell(row=max_row, column=c)
                cell.font = total_font
                
        # Auto-adjust column widths with some padding
        for c in range(1, max_col + 1):
            col_letter = openpyxl.utils.get_column_letter(c)
            max_len = 0
            for r in range(1, max_row + 1):
                val = ws.cell(row=r, column=c).value
                if val is not None:
                    max_len = max(max_len, len(str(val)))
            ws.column_dimensions[col_letter].width = max(max_len + 4, 12)
            
    wb.save(filename)
    print(f"Results exported to Excel: {filename}")

def export_results_csv(case, prefix, extra_results=None):
    case.update_generator_power()
    if hasattr(case, 'bus') and case.bus is not None:
        # Calculate generation per bus
        gen_p = np.zeros(len(case.bus))
        gen_q = np.zeros(len(case.bus))
        for i in range(len(case.gen)):
            if case.gen[i, GEN_STATUS] > 0:
                bus_idx = int(case.gen[i, GEN_BUS])
                gen_p[bus_idx] += case.gen[i, PG]
                gen_q[bus_idx] += case.gen[i, QG]

        bus_df = pd.DataFrame({
            "Bus_ID": case.external_bus_ids.astype(int),
            "TYPE": case.bus[:, BUS_TYPE].astype(int),
            "VM_pu": case.bus[:, VM],
            "VA_deg": case.bus[:, VA],
            "PD_MW": case.bus[:, PD],
            "QD_MVAr": case.bus[:, QD],
            "PG_MW": gen_p,
            "QG_MVAr": gen_q
        })
        bus_df.to_csv(f"{prefix}_bus.csv", index=False)
        
        gen_bus_ids = [int(case.external_bus_ids[int(b)]) for b in case.gen[:, GEN_BUS]]
        gen_df = pd.DataFrame({
            "Bus_ID": gen_bus_ids,
            "PG_MW": case.gen[:, PG],
            "QG_MVAr": case.gen[:, QG]
        })
        gen_df.to_csv(f"{prefix}_gen.csv", index=False)
        
        f_bus_ids = [int(case.external_bus_ids[int(b)]) for b in case.branch[:, F_BUS]]
        t_bus_ids = [int(case.external_bus_ids[int(b)]) for b in case.branch[:, T_BUS]]
        branch_df = pd.DataFrame({
            "From_Bus": f_bus_ids,
            "To_Bus": t_bus_ids,
            "BR_R": case.branch[:, BR_R],
            "BR_X": case.branch[:, BR_X],
            "BR_B": case.branch[:, BR_B],
            "RATE_A": case.branch[:, RATE_A],
            "PF_MW": case.branch[:, PF],
            "QF_MVAr": case.branch[:, QF],
            "PT_MW": case.branch[:, PT],
            "QT_MVAr": case.branch[:, QT]
        })
        branch_df.to_csv(f"{prefix}_branch.csv", index=False)

        # Line Flows & Losses CSV
        p_from = case.branch[:, PF]
        q_from = case.branch[:, QF]
        p_to = case.branch[:, PT]
        q_to = case.branch[:, QT]
        p_loss = p_from + p_to
        q_loss = q_from + q_to
        
        flows_df = pd.DataFrame({
            "From_Bus": f_bus_ids,
            "To_Bus": t_bus_ids,
            "P_FROM_TO_MW": p_from,
            "Q_FROM_TO_MVAR": q_from,
            "P_TO_FROM_MW": p_to,
            "Q_TO_FROM_MVAR": q_to,
            "P_LOSS_MW": p_loss,
            "Q_LOSS_MVAR": q_loss
        })
        
        total_row = {
            "From_Bus": "TOTAL",
            "To_Bus": "",
            "P_FROM_TO_MW": np.nan,
            "Q_FROM_TO_MVAR": np.nan,
            "P_TO_FROM_MW": np.nan,
            "Q_TO_FROM_MVAR": np.nan,
            "P_LOSS_MW": np.sum(p_loss),
            "Q_LOSS_MVAR": np.sum(q_loss)
        }
        flows_df = pd.concat([flows_df, pd.DataFrame([total_row])], ignore_index=True)
        flows_df.to_csv(f"{prefix}_flows.csv", index=False)
        
        print(f"Results exported to CSV: {prefix}_bus.csv, {prefix}_gen.csv, {prefix}_branch.csv, {prefix}_flows.csv")
        
    if hasattr(case, 'bus3p') and case.bus3p is not None:
        bus3p_df = pd.DataFrame(case.bus3p, columns=["Bus_ID", "Type", "BaseKV", "Vm1", "Vm2", "Vm3", "Va1", "Va2", "Va3"])
        bus3p_df.to_csv(f"{prefix}_bus3p.csv", index=False)
        print(f"Results exported to CSV: {prefix}_bus3p.csv")
        
    if extra_results is not None:
        if isinstance(extra_results, pd.DataFrame):
            extra_results.to_csv(f"{prefix}_extra.csv", index=False)
        elif isinstance(extra_results, dict):
            for k, v in extra_results.items():
                pd.DataFrame(v).to_csv(f"{prefix}_{k}.csv", index=False)
        print(f"Extra results exported to CSV.")

def export_results_html(case, filename, analysis, success, accuracy, extra_results=None, extra_info=None):
    # Calculate generation per bus
    gen_p = np.zeros(len(case.bus)) if (hasattr(case, 'bus') and case.bus is not None) else np.zeros(0)
    gen_q = np.zeros(len(case.bus)) if (hasattr(case, 'bus') and case.bus is not None) else np.zeros(0)
    if hasattr(case, 'gen') and case.gen is not None and len(gen_p) > 0:
        for i in range(len(case.gen)):
            if case.gen[i, GEN_STATUS] > 0:
                bus_idx = int(case.gen[i, GEN_BUS])
                if bus_idx < len(gen_p):
                    gen_p[bus_idx] += case.gen[i, PG]
                    gen_q[bus_idx] += case.gen[i, QG]

    analysis_upper = analysis.upper()
    status_text = "SUCCESS" if success else "FAILED"
    status_badge_class = "badge-success" if success else "badge-error"
    accuracy_str = f"{accuracy:.2e}"
    try:
        case_id = os.path.basename(filename).split('_', 1)[1].rsplit('.', 1)[0]
    except Exception:
        case_id = "Unknown"

    # Build Bus Table HTML
    if hasattr(case, 'bus3p') and case.bus3p is not None and len(case.bus3p) > 0:
        bus_rows_html = []
        for i in range(len(case.bus3p)):
            bus_id = int(case.bus3p[i, 0])
            vmag_a = case.bus3p[i, 3]
            vmag_b = case.bus3p[i, 4]
            vmag_c = case.bus3p[i, 5]
            vang_a = case.bus3p[i, 6]
            vang_b = case.bus3p[i, 7]
            vang_c = case.bus3p[i, 8]
            
            row = f"""<tr>
                <td class="text-center">{bus_id}</td>
                <td class="text-right">{vmag_a:.4f}</td>
                <td class="text-right">{vmag_b:.4f}</td>
                <td class="text-right">{vmag_c:.4f}</td>
                <td class="text-right">{vang_a:.2f}</td>
                <td class="text-right">{vang_b:.2f}</td>
                <td class="text-right">{vang_c:.2f}</td>
            </tr>"""
            bus_rows_html.append(row)
        bus_rows_str = "\n".join(bus_rows_html)
        bus_table_html = f"""<table id="bus-table">
            <thead>
                <tr>
                    <th scope="col" class="text-center">Bus ID</th>
                    <th scope="col" class="text-right">Va Mag (pu)</th>
                    <th scope="col" class="text-right">Vb Mag (pu)</th>
                    <th scope="col" class="text-right">Vc Mag (pu)</th>
                    <th scope="col" class="text-right">Va Ang (deg)</th>
                    <th scope="col" class="text-right">Vb Ang (deg)</th>
                    <th scope="col" class="text-right">Vc Ang (deg)</th>
                </tr>
            </thead>
            <tbody>
                {bus_rows_str}
            </tbody>
        </table>"""
    else:
        bus_rows_html = []
        limit = len(case.bus)
        for i in range(limit):
            bus_id = int(case.external_bus_ids[i])
            bus_t = int(case.bus[i, BUS_TYPE])
            type_str = "PQ" if bus_t == PQ else "PV" if bus_t == PV else "REF" if bus_t == REF else "Isolated"
            vm = case.bus[i, VM]
            va = case.bus[i, VA]
            pd_val = case.bus[i, PD]
            qd_val = case.bus[i, QD]
            gp = gen_p[i]
            gq = gen_q[i]
            
            row = f"""<tr>
                <td class="text-center">{bus_id}</td>
                <td class="text-center"><span class="badge" style="background-color: var(--bg-tertiary); color: var(--text-secondary);">{type_str}</span></td>
                <td class="text-right">{vm:.4f}</td>
                <td class="text-right">{va:.2f}</td>
                <td class="text-right">{pd_val:.2f}</td>
                <td class="text-right">{qd_val:.2f}</td>
                <td class="text-right">{gp:.2f}</td>
                <td class="text-right">{gq:.2f}</td>
            </tr>"""
            bus_rows_html.append(row)
        
        bus_rows_str = "\n".join(bus_rows_html)
        bus_table_html = f"""<table id="bus-table">
            <thead>
                <tr>
                    <th scope="col" class="text-center">Bus ID</th>
                    <th scope="col" class="text-center">Type</th>
                    <th scope="col" class="text-right">V Magnitude (pu)</th>
                    <th scope="col" class="text-right">V Angle (deg)</th>
                    <th scope="col" class="text-right">Load P (MW)</th>
                    <th scope="col" class="text-right">Load Q (MVAr)</th>
                    <th scope="col" class="text-right">Gen P (MW)</th>
                    <th scope="col" class="text-right">Gen Q (MVAr)</th>
                </tr>
            </thead>
            <tbody>
                {bus_rows_str}
            </tbody>
        </table>"""

    # Build Generator Table HTML
    gen_tab_btn = ""
    gen_table_content = ""
    if hasattr(case, 'gen') and case.gen is not None and len(case.gen) > 0:
        gen_tab_btn = """<button id="tab-btn-gen" class="tab-btn" onclick="switchTab('gen')" aria-selected="false" aria-controls="tab-content-gen">Generators</button>"""
        gen_rows_html = []
        for i in range(len(case.gen)):
            gen_id = i + 1
            bus_id = int(case.external_bus_ids[int(case.gen[i, GEN_BUS])])
            pg = case.gen[i, PG]
            qg = case.gen[i, QG]
            
            row = f"""<tr>
                <td class="text-center">{gen_id}</td>
                <td class="text-center">{bus_id}</td>
                <td class="text-right">{pg:.2f}</td>
                <td class="text-right">{qg:.2f}</td>
            </tr>"""
            gen_rows_html.append(row)
            
        gen_rows_str = "\n".join(gen_rows_html)
        gen_table_content = f"""<div id="tab-content-gen" class="tab-content" role="tabpanel" aria-labelledby="tab-btn-gen">
            <div class="table-container">
                <table id="gen-table">
                    <thead>
                        <tr>
                            <th scope="col" class="text-center">Gen ID</th>
                            <th scope="col" class="text-center">Bus ID</th>
                            <th scope="col" class="text-right">PG (MW)</th>
                            <th scope="col" class="text-right">QG (MVAr)</th>
                        </tr>
                    </thead>
                    <tbody>
                        {gen_rows_str}
                    </tbody>
                </table>
            </div>
        </div>"""

    # Build Line Flows Table HTML
    flows_tab_btn = ""
    flows_table_content = ""
    if hasattr(case, 'branch') and case.branch is not None and len(case.branch) > 0:
        flows_tab_btn = """<button id="tab-btn-flows" class="tab-btn" onclick="switchTab('flows')" aria-selected="false" aria-controls="tab-content-flows">Line Flows & Losses</button>"""
        flow_rows_html = []
        for i in range(len(case.branch)):
            f_bus = int(case.external_bus_ids[int(case.branch[i, F_BUS])])
            t_bus = int(case.external_bus_ids[int(case.branch[i, T_BUS])])
            pf = case.branch[i, PF]
            qf = case.branch[i, QF]
            pt = case.branch[i, PT]
            qt = case.branch[i, QT]
            pl = pf + pt
            ql = qf + qt
            
            row = f"""<tr>
                <td class="text-center">{f_bus}</td>
                <td class="text-center">{t_bus}</td>
                <td class="text-right">{pf:.2f}</td>
                <td class="text-right">{qf:.2f}</td>
                <td class="text-right">{pt:.2f}</td>
                <td class="text-right">{qt:.2f}</td>
                <td class="text-right">{pl:.2f}</td>
                <td class="text-right">{ql:.2f}</td>
            </tr>"""
            flow_rows_html.append(row)
            
        # Add Total Row
        full_p_loss = np.sum(case.branch[:, PF] + case.branch[:, PT])
        full_q_loss = np.sum(case.branch[:, QF] + case.branch[:, QT])
        total_row_html = f"""<tr class="total-row">
            <td class="text-center">Total</td>
            <td class="text-center"></td>
            <td class="text-right"></td>
            <td class="text-right"></td>
            <td class="text-right"></td>
            <td class="text-right"></td>
            <td class="text-right">{full_p_loss:.2f}</td>
            <td class="text-right">{full_q_loss:.2f}</td>
        </tr>"""
        flow_rows_html.append(total_row_html)
        
        flow_rows_str = "\n".join(flow_rows_html)
        flows_table_content = f"""<div id="tab-content-flows" class="tab-content" role="tabpanel" aria-labelledby="tab-btn-flows">
            <div class="table-container">
                <table id="flows-table">
                    <thead>
                        <tr>
                            <th scope="col" class="text-center">From Bus</th>
                            <th scope="col" class="text-center">To Bus</th>
                            <th scope="col" class="text-right">P From->To (MW)</th>
                            <th scope="col" class="text-right">Q From->To (MVAr)</th>
                            <th scope="col" class="text-right">P To->From (MW)</th>
                            <th scope="col" class="text-right">Q To->From (MVAr)</th>
                            <th scope="col" class="text-right">P Loss (MW)</th>
                            <th scope="col" class="text-right">Q Loss (MVAr)</th>
                        </tr>
                    </thead>
                    <tbody>
                        {flow_rows_str}
                    </tbody>
                </table>
            </div>
        </div>"""

    # Build System Quantities Card
    system_quantities_card = ""
    if hasattr(case, 'bus') and case.bus is not None and len(case.bus) > 0:
        vm = case.bus[:, VM]
        p_loss = np.sum(case.branch[:, PF] + case.branch[:, PT])
        q_loss = np.sum(case.branch[:, QF] + case.branch[:, QT])
        p_gen = np.sum(case.gen[:, PG]) if (hasattr(case, 'gen') and case.gen is not None) else 0.0
        
        system_quantities_card = f"""<div class="card">
            <h2>
                <svg width="20" height="20" fill="currentColor" viewBox="0 0 20 20"><path fill-rule="evenodd" d="M3 3a1 1 0 000 2v8a2 2 0 002 2h2.586l-1.293 1.293a1 1 0 101.414 1.414L9.414 15H10.586l1.707 1.707a1 1 0 001.414-1.414L12.414 15H15a2 2 0 002-2V5a1 1 0 00-2-2H3zm12 10H5V5h10v8z" clip-rule="evenodd"></path></svg>
                System Grid Metrics
            </h2>
            <ul class="metrics-list">
                <li>
                    <span class="label">Voltage Range</span>
                    <span class="value">Min: {np.min(vm):.4f} | Max: {np.max(vm):.4f} p.u.</span>
                </li>
                <li>
                    <span class="label">Active Power Loss</span>
                    <span class="value">{p_loss:.4f} MW</span>
                </li>
                <li>
                    <span class="label">Reactive Power Loss</span>
                    <span class="value">{q_loss:.4f} MVAr</span>
                </li>
                <li>
                    <span class="label">Total Dispatched Generation</span>
                    <span class="value">{p_gen:.2f} MW</span>
                </li>
            </ul>
        </div>"""

    # Build extra info metrics
    extra_summary_rows = ""
    if extra_info:
        rows_list = []
        for k, v in extra_info.items():
            rows_list.append(f"""<li>
                <span class="label">{k}</span>
                <span class="value">{v}</span>
            </li>""")
        extra_summary_rows = "\n".join(rows_list)

    # Build extra / optimization results tab
    extra_tab_btn = ""
    extra_table_content = ""
    if extra_results is not None:
        extra_tab_btn = """<button id="tab-btn-extra" class="tab-btn" onclick="switchTab('extra')" aria-selected="false" aria-controls="tab-content-extra">Optimization / Extra Outputs</button>"""
        
        if isinstance(extra_results, pd.DataFrame):
            headers_html = "".join([f'<th scope="col" class="text-center">{col}</th>' for col in extra_results.columns])
            rows_html = []
            for r_idx in range(min(20, len(extra_results))):
                cells_html = "".join([f'<td class="text-right">{extra_results.iloc[r_idx, c_idx]}</td>' for c_idx in range(len(extra_results.columns))])
                rows_html.append(f"<tr>{cells_html}</tr>")
            rows_str = "\n".join(rows_html)
            extra_table_content = f"""<div id="tab-content-extra" class="tab-content" role="tabpanel" aria-labelledby="tab-btn-extra">
                <div class="table-container">
                    <table id="extra-table">
                        <thead>
                            <tr>{headers_html}</tr>
                        </thead>
                        <tbody>
                            {rows_str}
                        </tbody>
                    </table>
                </div>
            </div>"""
        elif isinstance(extra_results, dict):
            dict_rows = []
            for k, v in extra_results.items():
                dict_rows.append(f"""<tr>
                    <td><strong>{k}</strong></td>
                    <td>{str(v)[:400]}</td>
                </tr>""")
            dict_rows_str = "\n".join(dict_rows)
            extra_table_content = f"""<div id="tab-content-extra" class="tab-content" role="tabpanel" aria-labelledby="tab-btn-extra">
                <div class="table-container">
                    <table id="extra-table-dict">
                        <thead>
                            <tr>
                                <th scope="col">Parameter</th>
                                <th scope="col">Value</th>
                            </tr>
                        </thead>
                        <tbody>
                            {dict_rows_str}
                        </tbody>
                    </table>
                </div>
            </div>"""

    # Semantic HTML template
    html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <meta name="description" content="Detailed power system simulation report generated by PowerPython.">
    <title>PowerPython Simulation Report - {analysis_upper}</title>
    <!-- Google Fonts -->
    <link rel="preconnect" href="https://fonts.googleapis.com">
    <link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
    <link href="https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;500;600;700;800&family=Plus+Jakarta+Sans:wght@300;400;500;600;700&display=swap" rel="stylesheet">
    <style>
        /* Modern CSS Reset & Variable Definitions */
        :root {{
            --bg-primary: #f8fafc;
            --bg-secondary: #ffffff;
            --bg-tertiary: #f1f5f9;
            --text-primary: #0f172a;
            --text-secondary: #475569;
            --text-muted: #64748b;
            --accent: #0f766e;
            --accent-light: #ccfbf1;
            --accent-hover: #115e59;
            --success: #15803d;
            --success-light: #dcfce7;
            --error: #b91c1c;
            --error-light: #fee2e2;
            --border: #e2e8f0;
            --shadow: 0 4px 6px -1px rgb(0 0 0 / 0.05), 0 2px 4px -2px rgb(0 0 0 / 0.05);
            --shadow-lg: 0 10px 15px -3px rgb(0 0 0 / 0.05), 0 4px 6px -4px rgb(0 0 0 / 0.05);
            --transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
            --font-main: 'Plus Jakarta Sans', sans-serif;
            --font-heading: 'Outfit', sans-serif;
        }}

        [data-theme="dark"] {{
            --bg-primary: #0f172a;
            --bg-secondary: #1e293b;
            --bg-tertiary: #334155;
            --text-primary: #f8fafc;
            --text-secondary: #cbd5e1;
            --text-muted: #94a3b8;
            --accent: #2dd4bf;
            --accent-light: #115e59;
            --accent-hover: #5eead4;
            --success: #4ade80;
            --success-light: #14532d;
            --error: #f87171;
            --error-light: #7f1d1d;
            --border: #334155;
            --shadow: 0 4px 6px -1px rgb(0 0 0 / 0.3), 0 2px 4px -2px rgb(0 0 0 / 0.3);
            --shadow-lg: 0 10px 15px -3px rgb(0 0 0 / 0.4), 0 4px 6px -4px rgb(0 0 0 / 0.4);
        }}

        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}

        body {{
            background-color: var(--bg-primary);
            color: var(--text-primary);
            font-family: var(--font-main);
            line-height: 1.5;
            transition: var(--transition);
            padding: 2rem 1rem;
        }}

        .container {{
            max-width: 1200px;
            margin: 0 auto;
        }}

        /* Header section with Glassmorphism */
        header {{
            background: rgba(255, 255, 255, 0.8);
            backdrop-filter: blur(12px);
            -webkit-backdrop-filter: blur(12px);
            border: 1px solid rgba(255, 255, 255, 0.3);
            border-radius: 1.5rem;
            padding: 2rem;
            margin-bottom: 2rem;
            display: flex;
            justify-content: space-between;
            align-items: center;
            box-shadow: var(--shadow-lg);
            transition: var(--transition);
        }}

        [data-theme="dark"] header {{
            background: rgba(30, 41, 59, 0.8);
            border: 1px solid rgba(255, 255, 255, 0.05);
        }}

        .header-title h1 {{
            font-family: var(--font-heading);
            font-size: 2.25rem;
            font-weight: 800;
            background: linear-gradient(135deg, var(--accent), #3b82f6);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
            margin-bottom: 0.25rem;
        }}

        .header-title p {{
            color: var(--text-muted);
            font-size: 0.95rem;
            font-weight: 500;
        }}

        /* Theme Toggle Switch */
        .theme-toggle {{
            background: var(--bg-tertiary);
            border: 1px solid var(--border);
            padding: 0.5rem 1rem;
            border-radius: 9999px;
            cursor: pointer;
            display: flex;
            align-items: center;
            gap: 0.5rem;
            font-family: var(--font-main);
            font-size: 0.875rem;
            font-weight: 600;
            color: var(--text-secondary);
            transition: var(--transition);
        }}

        .theme-toggle:hover {{
            background: var(--border);
            color: var(--text-primary);
        }}

        /* Cards layout */
        .grid {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(300px, 1fr));
            gap: 1.5rem;
            margin-bottom: 2rem;
        }}

        .card {{
            background-color: var(--bg-secondary);
            border: 1px solid var(--border);
            border-radius: 1.25rem;
            padding: 1.5rem;
            box-shadow: var(--shadow);
            transition: var(--transition);
        }}

        .card:hover {{
            transform: translateY(-2px);
            box-shadow: var(--shadow-lg);
        }}

        .card h2 {{
            font-family: var(--font-heading);
            font-size: 1.25rem;
            font-weight: 700;
            margin-bottom: 1rem;
            display: flex;
            align-items: center;
            gap: 0.5rem;
            color: var(--text-primary);
        }}

        .metrics-list {{
            list-style: none;
        }}

        .metrics-list li {{
            display: flex;
            justify-content: space-between;
            align-items: center;
            padding: 0.75rem 0;
            border-bottom: 1px solid var(--border);
        }}

        .metrics-list li:last-child {{
            border-bottom: none;
        }}

        .metrics-list .label {{
            color: var(--text-secondary);
            font-size: 0.9rem;
            font-weight: 500;
        }}

        .metrics-list .value {{
            font-weight: 700;
            font-size: 1rem;
        }}

        /* Status badges */
        .badge {{
            padding: 0.25rem 0.75rem;
            border-radius: 9999px;
            font-size: 0.75rem;
            font-weight: 700;
            text-transform: uppercase;
        }}

        .badge-success {{
            background-color: var(--success-light);
            color: var(--success);
        }}

        .badge-error {{
            background-color: var(--error-light);
            color: var(--error);
        }}

        /* Navigation Tabs */
        .tabs {{
            display: flex;
            gap: 0.5rem;
            border-bottom: 2px solid var(--border);
            margin-bottom: 1.5rem;
            overflow-x: auto;
            padding-bottom: 2px;
        }}

        .tab-btn {{
            background: none;
            border: none;
            color: var(--text-muted);
            font-family: var(--font-main);
            font-size: 0.95rem;
            font-weight: 600;
            padding: 0.75rem 1.25rem;
            cursor: pointer;
            border-radius: 0.5rem 0.5rem 0 0;
            transition: var(--transition);
            white-space: nowrap;
        }}

        .tab-btn:hover {{
            color: var(--text-primary);
            background-color: var(--bg-tertiary);
        }}

        .tab-btn.active {{
            color: var(--accent);
            border-bottom: 3px solid var(--accent);
            margin-bottom: -2px;
        }}

        /* Table Design (Premium, WCAG 2.2 Compliant) */
        .tab-content {{
            display: none;
            animation: fadeIn 0.4s ease;
        }}

        .tab-content.active {{
            display: block;
        }}

        @keyframes fadeIn {{
            from {{ opacity: 0; transform: translateY(4px); }}
            to {{ opacity: 1; transform: translateY(0); }}
        }}

        .table-container {{
            width: 100%;
            overflow-x: auto;
            border-radius: 1rem;
            border: 1px solid var(--border);
            box-shadow: var(--shadow);
            background-color: var(--bg-secondary);
        }}

        table {{
            width: 100%;
            border-collapse: collapse;
            text-align: left;
            font-size: 0.9rem;
        }}

        thead {{
            background-color: var(--bg-tertiary);
        }}

        th {{
            padding: 1rem 1.25rem;
            font-family: var(--font-heading);
            font-weight: 700;
            color: var(--text-secondary);
            font-size: 0.85rem;
            text-transform: uppercase;
            letter-spacing: 0.05em;
            border-bottom: 2px solid var(--border);
            white-space: nowrap;
        }}

        td {{
            padding: 0.875rem 1.25rem;
            border-bottom: 1px solid var(--border);
            color: var(--text-secondary);
            font-weight: 500;
        }}

        tbody tr:hover {{
            background-color: var(--bg-primary);
            transition: var(--transition);
        }}

        tbody tr:last-child td {{
            border-bottom: none;
        }}

        /* Total row special styling */
        tr.total-row {{
            background-color: var(--bg-tertiary);
            font-weight: 700;
        }}

        tr.total-row td {{
            color: var(--text-primary);
            font-weight: 700;
            border-top: 2px solid var(--border);
            border-bottom: 2px solid var(--border);
        }}

        /* Numeric alignment */
        .text-right {{
            text-align: right;
        }}

        .text-center {{
            text-align: center;
        }}

        /* Footer styling */
        footer {{
            margin-top: 4rem;
            text-align: center;
            color: var(--text-muted);
            font-size: 0.8rem;
            font-weight: 500;
            border-top: 1px solid var(--border);
            padding-top: 2rem;
        }}

        footer a {{
            color: var(--accent);
            text-decoration: none;
            font-weight: 600;
        }}

        footer a:hover {{
            text-decoration: underline;
        }}

        /* Responsive adjustments */
        @media (max-width: 768px) {{
            body {{
                padding: 1rem 0.5rem;
            }}
            header {{
                flex-direction: column;
                align-items: flex-start;
                gap: 1rem;
            }}
            .theme-toggle {{
                align-self: flex-end;
            }}
        }}
    </style>
</head>
<body data-theme="light">
    <div class="container">
        <!-- HEADER -->
        <header>
            <div class="header-title">
                <h1>PowerPython Simulation Report</h1>
                <p>Analysis: {analysis_upper} | Target Case: {case_id}</p>
            </div>
            <button id="themeToggleBtn" class="theme-toggle" aria-label="Toggle visual theme" onclick="toggleTheme()">
                <span id="themeToggleText">Dark Mode</span>
            </button>
        </header>

        <!-- MAIN LAYOUT -->
        <main>
            <!-- OVERVIEW SUMMARY CARDS -->
            <section aria-labelledby="summary-heading" class="grid">
                <!-- Card 1: Convergence Status -->
                <div class="card">
                    <h2 id="summary-heading">
                        <svg width="20" height="20" fill="currentColor" viewBox="0 0 20 20"><path fill-rule="evenodd" d="M10 18a8 8 0 100-16 8 8 0 000 16zm3.707-9.293a1 1 0 00-1.414-1.414L9 10.586 7.707 9.293a1 1 0 00-1.414 1.414l2 2a1 1 0 001.414 0l4-4z" clip-rule="evenodd"></path></svg>
                        Convergence Summary
                    </h2>
                    <ul class="metrics-list">
                        <li>
                            <span class="label">Status</span>
                            <span class="value"><span class="badge {status_badge_class}">{status_text}</span></span>
                        </li>
                        <li>
                            <span class="label">Tolerance/Accuracy</span>
                            <span class="value">{accuracy_str}</span>
                        </li>
                        {extra_summary_rows}
                    </ul>
                </div>

                <!-- Card 2: System Quantities -->
                {system_quantities_card}
            </section>

            <!-- TABS NAVIGATION -->
            <nav aria-label="Report section navigation">
                <div class="tabs">
                    <button id="tab-btn-bus" class="tab-btn active" onclick="switchTab('bus')" aria-selected="true" aria-controls="tab-content-bus">Buses & Voltages</button>
                    {gen_tab_btn}
                    {flows_tab_btn}
                    {extra_tab_btn}
                </div>
            </nav>

            <!-- TAB CONTENTS -->
            <!-- Tab 1: Buses -->
            <div id="tab-content-bus" class="tab-content active" role="tabpanel" aria-labelledby="tab-btn-bus">
                <div class="table-container">
                    {bus_table_html}
                </div>
            </div>

            <!-- Tab 2: Generators -->
            {gen_table_content}

            <!-- Tab 3: Line Flows -->
            {flows_table_content}

            <!-- Tab 4: Extras -->
            {extra_table_content}
        </main>

        <!-- FOOTER -->
        <footer>
            <p>Generated by <a href="https://github.com/salorajan/matpower_data_migration" target="_blank" rel="noopener">PowerPython (matpower-python)</a> &copy; 2026. All rights reserved.</p>
        </footer>
    </div>

    <!-- JAVASCRIPT FOR DYNAMIC INTERACTION -->
    <script>
        // Theme switching logic
        function toggleTheme() {{
            const body = document.body;
            const currentTheme = body.getAttribute('data-theme');
            const newTheme = currentTheme === 'dark' ? 'light' : 'dark';
            body.setAttribute('data-theme', newTheme);
            
            const btnText = document.getElementById('themeToggleText');
            btnText.textContent = newTheme === 'dark' ? 'Light Mode' : 'Dark Mode';
            localStorage.setItem('power_python_theme', newTheme);
        }}

        // Initialize theme from storage
        const savedTheme = localStorage.getItem('power_python_theme');
        if (savedTheme) {{
            document.body.setAttribute('data-theme', savedTheme);
            document.getElementById('themeToggleText').textContent = savedTheme === 'dark' ? 'Light Mode' : 'Dark Mode';
        }}

        // Tabs switching logic
        function switchTab(tabId) {{
            // Deactivate all tabs and tab content
            const buttons = document.querySelectorAll('.tab-btn');
            const contents = document.querySelectorAll('.tab-content');
            
            buttons.forEach(btn => {{
                btn.classList.remove('active');
                btn.setAttribute('aria-selected', 'false');
            }});
            
            contents.forEach(content => {{
                content.classList.remove('active');
            }});
            
            // Activate selected tab and content
            const targetBtn = document.getElementById('tab-btn-' + tabId);
            const targetContent = document.getElementById('tab-content-' + tabId);
            
            if (targetBtn && targetContent) {{
                targetBtn.classList.add('active');
                targetBtn.setAttribute('aria-selected', 'true');
                targetContent.classList.add('active');
            }}
        }}
    </script>
</body>
</html>
"""
    with open(filename, 'w', encoding='utf-8') as f:
        f.write(html_content)
    print(f"Results exported to HTML: {filename}")

def make_table_wcag_compliant(table, title, description):
    """
    Apply WCAG 2.2 compatibility features to a python-docx table:
    1. Set the first row (headers) to repeat on every page (tblHeader).
    2. Prevent all rows from splitting across pages (cantSplit).
    3. Add alternative text (Title and Description) for screen readers.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    # 1. Set repeat header for row 0
    if len(table.rows) > 0:
        tr = table.rows[0]._tr
        trPr = tr.get_or_add_trPr()
        tblHeader = OxmlElement('w:tblHeader')
        tblHeader.set(qn('w:val'), "true")
        trPr.append(tblHeader)
        
    # 2. Set cantSplit for all rows
    for row in table.rows:
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        cantSplit = OxmlElement('w:cantSplit')
        trPr.append(cantSplit)
        
    # 3. Add accessibility Title and Description (Alt text)
    tblPr = table._tbl.tblPr
    tblCaption = OxmlElement('w:tblCaption')
    tblCaption.set(qn('w:val'), title)
    tblPr.append(tblCaption)
    
    tblDescription = OxmlElement('w:tblDescription')
    tblDescription.set(qn('w:val'), description)
    tblPr.append(tblDescription)

def export_results_docx(case, filename, analysis, success, accuracy, extra_results=None, extra_info=None):
    doc = docx.Document()
    
    # Header
    title = doc.add_paragraph()
    run = title.add_run("PowerPython Simulation Report")
    run.font.size = Pt(24)
    run.font.bold = True
    
    # Subtitle
    sub = doc.add_paragraph()
    grid_size = case.external_bus_ids.size if hasattr(case, 'external_bus_ids') else len(case.bus3p)
    run_sub = sub.add_run(f"Method: {analysis.upper()} | Case Size: {grid_size}-Bus Grid")
    run_sub.font.size = Pt(14)
    run_sub.italic = True
    
    # Section 1: Summary
    doc.add_heading("1.0 Simulation Summary", level=1)
    
    p = doc.add_paragraph()
    p.add_run("Convergence Status: ").bold = True
    p.add_run("SUCCESS\n" if success else "FAILED\n")
    p.add_run("Accuracy/Tolerance: ").bold = True
    p.add_run(f"{accuracy:.2e}\n")
    
    if hasattr(case, 'bus') and case.bus is not None:
        vm = case.bus[:, VM]
        p.add_run("Voltage Range: ").bold = True
        p.add_run(f"Min: {np.min(vm):.4f} p.u. | Max: {np.max(vm):.4f} p.u.\n")
        
        p_loss = np.sum(case.branch[:, PF] + case.branch[:, PT])
        q_loss = np.sum(case.branch[:, QF] + case.branch[:, QT])
        p.add_run("System Losses: ").bold = True
        p.add_run(f"P_loss: {p_loss:.4f} MW | Q_loss: {q_loss:.4f} MVAr\n")
        
        p_gen = np.sum(case.gen[:, PG])
        p.add_run("Total Generation: ").bold = True
        p.add_run(f"P_gen: {p_gen:.2f} MW\n")
        
    if extra_info:
        for k, v in extra_info.items():
            p.add_run(f"{k}: ").bold = True
            p.add_run(f"{v}\n")
            
    # Section 2: Tables
    if hasattr(case, 'bus') and case.bus is not None:
        # Calculate generation per bus
        gen_p = np.zeros(len(case.bus))
        gen_q = np.zeros(len(case.bus))
        for i in range(len(case.gen)):
            if case.gen[i, GEN_STATUS] > 0:
                bus_idx = int(case.gen[i, GEN_BUS])
                gen_p[bus_idx] += case.gen[i, PG]
                gen_q[bus_idx] += case.gen[i, QG]
                
        doc.add_heading("2.0 Bus Voltages and Power Dispatch", level=1)
        table = doc.add_table(rows=1, cols=8)
        table.style = 'Light Shading Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Bus ID'
        hdr_cells[1].text = 'Type'
        hdr_cells[2].text = 'V Magnitude (pu)'
        hdr_cells[3].text = 'V Angle (deg)'
        hdr_cells[4].text = 'Load P (MW)'
        hdr_cells[5].text = 'Load Q (MVAr)'
        hdr_cells[6].text = 'Gen P (MW)'
        hdr_cells[7].text = 'Gen Q (MVAr)'
        
        limit = min(100, len(case.bus))
        for i in range(limit):
            row_cells = table.add_row().cells
            row_cells[0].text = str(int(case.external_bus_ids[i]))
            
            # Map type to string: PQ, PV, REF
            bus_t = int(case.bus[i, BUS_TYPE])
            type_str = "PQ" if bus_t == PQ else "PV" if bus_t == PV else "REF" if bus_t == REF else "Isolated"
            row_cells[1].text = type_str
            
            row_cells[2].text = f"{case.bus[i, VM]:.4f}"
            row_cells[3].text = f"{case.bus[i, VA]:.2f}"
            row_cells[4].text = f"{case.bus[i, PD]:.2f}"
            row_cells[5].text = f"{case.bus[i, QD]:.2f}"
            row_cells[6].text = f"{gen_p[i]:.2f}"
            row_cells[7].text = f"{gen_q[i]:.2f}"
            
        make_table_wcag_compliant(table, "Bus Voltages and Power Dispatch Table", 
                                  "Table showing voltage magnitude, angle, and active/reactive load and generation for each bus.")
            
        doc.add_heading("3.0 Generator Dispatch", level=1)
        table_gen = doc.add_table(rows=1, cols=4)
        table_gen.style = 'Light Shading Accent 1'
        hdr_gen = table_gen.rows[0].cells
        hdr_gen[0].text = 'Gen ID'
        hdr_gen[1].text = 'Bus ID'
        hdr_gen[2].text = 'PG (MW)'
        hdr_gen[3].text = 'QG (MVAr)'
        
        for i in range(len(case.gen)):
            row_cells = table_gen.add_row().cells
            row_cells[0].text = str(i + 1)
            row_cells[1].text = str(int(case.external_bus_ids[int(case.gen[i, GEN_BUS])]))
            row_cells[2].text = f"{case.gen[i, PG]:.2f}"
            row_cells[3].text = f"{case.gen[i, QG]:.2f}"
            
        make_table_wcag_compliant(table_gen, "Generator Dispatch Table", 
                                  "Table listing generator indices, connection buses, and real and reactive power outputs.")
            
        doc.add_heading("4.0 Line Flows and Losses", level=1)
        table_flows = doc.add_table(rows=1, cols=8)
        table_flows.style = 'Light Shading Accent 1'
        hdr_flows = table_flows.rows[0].cells
        hdr_flows[0].text = 'From Bus'
        hdr_flows[1].text = 'To Bus'
        hdr_flows[2].text = 'P From->To (MW)'
        hdr_flows[3].text = 'Q From->To (MVAr)'
        hdr_flows[4].text = 'P To->From (MW)'
        hdr_flows[5].text = 'Q To->From (MVAr)'
        hdr_flows[6].text = 'P Loss (MW)'
        hdr_flows[7].text = 'Q Loss (MVAr)'
        
        limit_br = min(100, len(case.branch))
        for i in range(limit_br):
            row_cells = table_flows.add_row().cells
            f_id = int(case.external_bus_ids[int(case.branch[i, F_BUS])])
            t_id = int(case.external_bus_ids[int(case.branch[i, T_BUS])])
            pf = case.branch[i, PF]
            qf = case.branch[i, QF]
            pt = case.branch[i, PT]
            qt = case.branch[i, QT]
            pl = pf + pt
            ql = qf + qt
            
            row_cells[0].text = str(f_id)
            row_cells[1].text = str(t_id)
            row_cells[2].text = f"{pf:.2f}"
            row_cells[3].text = f"{qf:.2f}"
            row_cells[4].text = f"{pt:.2f}"
            row_cells[5].text = f"{qt:.2f}"
            row_cells[6].text = f"{pl:.2f}"
            row_cells[7].text = f"{ql:.2f}"
            
        # Add Total Row
        row_cells = table_flows.add_row().cells
        row_cells[0].text = "Total"
        row_cells[1].text = ""
        row_cells[2].text = ""
        row_cells[3].text = ""
        row_cells[4].text = ""
        row_cells[5].text = ""
        full_p_loss = np.sum(case.branch[:, PF] + case.branch[:, PT])
        full_q_loss = np.sum(case.branch[:, QF] + case.branch[:, QT])
        row_cells[6].text = f"{full_p_loss:.2f}"
        row_cells[7].text = f"{full_q_loss:.2f}"
        
        # Bold the total row
        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    
        make_table_wcag_compliant(table_flows, "Line Flows and Losses Table", 
                                  "Table showing active and reactive power flows in both directions, and individual line active/reactive losses with totals.")
            
    elif hasattr(case, 'bus3p') and case.bus3p is not None:
        doc.add_heading("2.0 3-Phase Bus Voltages", level=1)
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Light Shading Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Bus ID'
        hdr_cells[1].text = 'Va (pu)'
        hdr_cells[2].text = 'Vb (pu)'
        hdr_cells[3].text = 'Vc (pu)'
        hdr_cells[4].text = 'AngA'
        hdr_cells[5].text = 'AngB'
        hdr_cells[6].text = 'AngC'
        
        for i in range(len(case.bus3p)):
            row_cells = table.add_row().cells
            row_cells[0].text = str(int(case.bus3p[i, 0]))
            row_cells[1].text = f"{case.bus3p[i, 3]:.4f}"
            row_cells[2].text = f"{case.bus3p[i, 4]:.4f}"
            row_cells[3].text = f"{case.bus3p[i, 5]:.4f}"
            row_cells[4].text = f"{case.bus3p[i, 6]:.2f}"
            row_cells[5].text = f"{case.bus3p[i, 7]:.2f}"
            row_cells[6].text = f"{case.bus3p[i, 8]:.2f}"
            
        make_table_wcag_compliant(table, "3-Phase Bus Voltages Table", 
                                  "Table listing voltage magnitudes and angles for phases A, B, and C at each bus.")
            
    # Section 5: Extra Results if present
    if extra_results is not None:
        doc.add_heading("5.0 Optimization / Extra Outputs", level=1)
        if isinstance(extra_results, pd.DataFrame):
            # Render first 20 rows of DataFrame
            table_extra = doc.add_table(rows=1, cols=len(extra_results.columns))
            table_extra.style = 'Light Shading Accent 1'
            for col_idx, col_name in enumerate(extra_results.columns):
                table_extra.rows[0].cells[col_idx].text = str(col_name)
            for r_idx in range(min(20, len(extra_results))):
                row_cells = table_extra.add_row().cells
                for col_idx, col_name in enumerate(extra_results.columns):
                    row_cells[col_idx].text = f"{extra_results.iloc[r_idx, col_idx]}"
            make_table_wcag_compliant(table_extra, "Optimization Extra Results Table", 
                                      "Table displaying extra optimization results and outputs.")
        elif isinstance(extra_results, dict):
            for k, v in extra_results.items():
                doc.add_paragraph().add_run(f"Data: {k}").bold = True
                p_dict = doc.add_paragraph()
                p_dict.add_run(str(v)[:800])
                
    doc.save(filename)
    print(f"Results exported to Word Document: {filename}")

def print_help_console(analysis):
    help_text = HELP_DICT.get(analysis, f"No detailed help available for {analysis}.")
    print("\n" + "="*80)
    print(f"{analysis.upper()} SOLVER HELP")
    print("="*80)
    print(help_text)
    print("="*80 + "\n")

def export_help_docx(analysis):
    help_text = HELP_DICT.get(analysis, f"No detailed help available for {analysis}.")
    filename = f"{analysis}_help.docx"
    doc = docx.Document()
    doc.add_heading(f"PowerPython Solver Help: {analysis.upper()}", level=0)
    
    for line in help_text.split('\n'):
        if line.startswith("Usage:"):
            doc.add_heading("Usage", level=1)
            doc.add_paragraph(line)
        elif line.startswith("Example:"):
            p = doc.add_paragraph()
            p.add_run("Example: ").bold = True
            p.add_run(line[8:])
        elif line.startswith("Physics/Algorithm:"):
            doc.add_heading("Physics & Algorithm", level=1)
        else:
            doc.add_paragraph(line)
            
    doc.save(filename)
    print(f"Help content exported to Word Document: {filename}")

def run_cli_command(analysis, args):
    is_help = any(h in args for h in ["--help", "-h", "help"])
    if is_help:
        if "docx" in args:
            export_help_docx(analysis)
        else:
            print_help_console(analysis)
        return
        
    if len(args) < 1:
        print(f"Error: Missing case_id parameter.")
        print_help_console(analysis)
        return
        
    case_id = args[0]
    
    # Parse accuracy dynamically (e.g. 1e-4, 0.0001, 1.0e-5)
    accuracy = 1e-8
    for arg in args[1:]:
        try:
            if ("e" in arg or "." in arg or arg.isdigit()) and not any(ext in arg.lower() for ext in ["excel", "xlsx", "csv", "docx", "word", "html", "htm"]):
                accuracy = float(arg)
                break
        except ValueError:
            pass
            
    # Parse export format
    export_format = None
    for arg in args[1:]:
        arg_lower = arg.lower()
        if arg_lower in ["excel", "xlsx"]:
            export_format = "excel"
            break
        elif arg_lower in ["csv"]:
            export_format = "csv"
            break
        elif arg_lower in ["docx", "word"]:
            export_format = "docx"
            break
        elif arg_lower in ["html", "htm"]:
            export_format = "html"
            break

    # Load case
    case = load_case_by_id(case_id)
    if case is None:
        return
        
    print(f"Running {analysis.upper()} on Case {case_id} (accuracy={accuracy:.2e})...")
    
    success = False
    extra_results = None
    extra_info = {}
    
    if analysis == 'acpf':
        case, success = run_power_flow(case, algorithm='nr', tol=accuracy, verbose=True)
    elif analysis == 'gausspf':
        case, success = run_power_flow(case, algorithm='gs', tol=accuracy, verbose=True)
    elif analysis == 'fdpf':
        case, success = run_power_flow(case, algorithm='fd', tol=accuracy, verbose=True)
    elif analysis == 'dcpf':
        case, success = run_dc_pf(case)
    elif analysis == 'hepf':
        # HEPF check accuracy dynamically by matching iterations
        case, success = run_hepf(case, verbose=True)
    elif analysis == 'cnr':
        case, success = run_complex_nr(case, tol=accuracy, verbose=True)
    elif analysis == 'pf3p':
        case, success = run_3p_pf(case, tol=accuracy, verbose=True)
    elif analysis == 'radial':
        case, success = run_radial_pf(case, tol=accuracy, verbose=True)
    elif analysis == 'dcopf':
        case, success = run_dc_opf(case, verbose=True)
        if success:
            extra_info["Total Optimized Cost"] = f"${calculate_total_cost(case):,.2f}"
    elif analysis == 'acopf':
        case, success = run_ac_opf(case, verbose=True)
        if success:
            extra_info["Total Optimized Cost"] = f"${calculate_total_cost(case):,.2f}"
    elif analysis == 'sdpopf':
        case, success = run_sdp_opf(case, verbose=True)
        if success:
            extra_info["Total Lower Bound Cost"] = f"${calculate_total_cost(case):,.2f}"
    elif analysis == 'uopf':
        case, success = run_uopf(case, solver='dcopf', verbose=True)
        if success:
            extra_info["Optimized Cost"] = f"${calculate_total_cost(case):,.2f}"
            extra_info["Generators Online"] = f"{np.sum(case.gen[:, GEN_STATUS] > 0)} / {len(case.gen)}"
    elif analysis == 'scopf':
        case, success = run_sc_opf(case, verbose=True)
    elif analysis == 'opf3p':
        case, success = run_3p_opf(case, verbose=True)
    elif analysis == 'mpopf':
        nt = 6
        if not case.profiles.get('load'): case.profiles['load'] = [1.0, 1.2, 1.5, 1.3, 1.1, 0.9]
        if not case.storage: case.storage = {'idx': [0], 'MaxCharge': [10.0], 'MaxDischarge': [10.0], 'InEff': [0.95], 'OutEff': [0.95], 'MinSOC': [0.0], 'MaxSOC': [50.0], 'InitialSOC': [25.0]}
        results, success = run_mp_opf(case, nt=nt, verbose=True)
        if success:
            ng = len(case.gen)
            extra_results = pd.DataFrame(results['Pg'], columns=[f"Gen {i+1}" for i in range(ng)])
            extra_info["Total Operating Cost"] = f"${results['Cost']:,.2f}"
    elif analysis == 'stopf':
        results, success = run_stochastic_opf(case, verbose=True)
        if success:
            extra_results = pd.DataFrame(results['Pg'])
            extra_info["Expected Operating Cost"] = f"${results['Expected_Cost']:,.2f}"
    elif analysis == 'market':
        ng = len(case.gen)
        offers = {'qty': [[50, 50, 50]] * ng, 'prc': [[20, 40, 60]] * ng}
        case, mkt_df = run_market_auction(case, offers, verbose=True)
        if mkt_df is not None:
            success = True
            extra_results = mkt_df
            extra_info["Total Market Turnover"] = f"${mkt_df['Revenue_$'].sum():,.2f}"
    elif analysis == 'varplan':
        case, success = run_var_planning(case, verbose=True)
        if success:
            extra_info["Total Compensation Added"] = f"{np.sum(case.bus[:, BS]):.2f} MVAr"
    elif analysis == 'contingency':
        df = run_contingency_analysis(case, verbose=True)
        if not df.empty:
            success = True
            extra_results = df
            extra_info["Violations Detected"] = f"{len(df)} outages exceeded limits"
        else:
            success = True
            print("No line flow violations detected.")
    elif analysis == 'se':
        case, success = run_state_estimation(case, verbose=True)
    elif analysis == 'cpf':
        results = run_cpf(case, verbose=True)
        if results:
            success = True
            extra_results = pd.DataFrame(results, columns=['lambda', 'vm_avg'])
            extra_info["Maximum System Loadability (Lambda)"] = f"{results[-1][0]:.4f}"
    elif analysis == 'audit':
        case.to_internal()
        balance = calculate_system_balance(case)
        success = True
        extra_info["P Residual"] = f"{balance['residual_p']:.4f} MW"
        extra_info["Q Residual"] = f"{balance['residual_q']:.4f} MVAr"
        print(f"Audit Residual P: {balance['residual_p']:.4f} MW | Q: {balance['residual_q']:.4f} MVAr")
    elif analysis == 'lmp':
        case, success = run_dc_opf(case, verbose=False)
        if success:
            df = decompose_dc_lmp(case)
            extra_results = df
            extra_info["Reference Energy Price"] = f"${df['Energy'].iloc[0]:.2f} / MWh"
            extra_info["Max Congestion component"] = f"${df['Congestion'].max():.2f} / MWh"
            
    # Export results
    if success:
        if export_format == "excel":
            export_results_excel(case, f"{analysis}_{case_id}.xlsx", analysis, extra_results)
        elif export_format == "csv":
            export_results_csv(case, f"{analysis}_{case_id}", extra_results)
        elif export_format == "docx":
            export_results_docx(case, f"{analysis}_{case_id}.docx", analysis, success, accuracy, extra_results, extra_info)
        elif export_format == "html":
            export_results_html(case, f"{analysis}_{case_id}.html", analysis, success, accuracy, extra_results, extra_info)
        else:
            # Print basic report
            if hasattr(case, 'bus') and case.bus is not None and len(case.bus) > 0:
                vm = case.bus[:, VM]
                print(f"\n{analysis.upper()} Success! Voltages: Min={np.min(vm):.4f} pu, Max={np.max(vm):.4f} pu")
            elif hasattr(case, 'bus3p') and case.bus3p is not None and len(case.bus3p) > 0:
                vm = case.bus3p[:, 3:6]
                print(f"\n{analysis.upper()} Success! Phase Voltages: Min={np.min(vm):.4f} pu, Max={np.max(vm):.4f} pu")
    else:
        print(f"Simulation failed or did not converge for case {case_id}")

# 23 Console Entry Points
def cli_acpf(): run_cli_command('acpf', sys.argv[1:])
def cli_gausspf(): run_cli_command('gausspf', sys.argv[1:])
def cli_fdpf(): run_cli_command('fdpf', sys.argv[1:])
def cli_dcpf(): run_cli_command('dcpf', sys.argv[1:])
def cli_hepf(): run_cli_command('hepf', sys.argv[1:])
def cli_cnr(): run_cli_command('cnr', sys.argv[1:])
def cli_pf3p(): run_cli_command('pf3p', sys.argv[1:])
def cli_radial(): run_cli_command('radial', sys.argv[1:])
def cli_dcopf(): run_cli_command('dcopf', sys.argv[1:])
def cli_acopf(): run_cli_command('acopf', sys.argv[1:])
def cli_sdpopf(): run_cli_command('sdpopf', sys.argv[1:])
def cli_uopf(): run_cli_command('uopf', sys.argv[1:])
def cli_scopf(): run_cli_command('scopf', sys.argv[1:])
def cli_opf3p(): run_cli_command('opf3p', sys.argv[1:])
def cli_mpopf(): run_cli_command('mpopf', sys.argv[1:])
def cli_stopf(): run_cli_command('stopf', sys.argv[1:])
def cli_market(): run_cli_command('market', sys.argv[1:])
def cli_varplan(): run_cli_command('varplan', sys.argv[1:])
def cli_contingency(): run_cli_command('contingency', sys.argv[1:])
def cli_se(): run_cli_command('se', sys.argv[1:])
def cli_cpf(): run_cli_command('cpf', sys.argv[1:])
def cli_audit(): run_cli_command('audit', sys.argv[1:])
def cli_lmp(): run_cli_command('lmp', sys.argv[1:])

def main():
    if len(sys.argv) < 3:
        print("\n" + "="*80)
        print(f"{'POWERPYTHON UNIFIED COMMAND LINE INTERFACE':^80}")
        print("="*80)
        print("\nUsage: power-cli <analysis> <case_id> [accuracy] [export_format]")
        print("Or run directly: <analysis> <case_id> [accuracy] [export_format]")
        print("\nAVAILABLE COMMANDS:")
        print(f"  Power Flows : acpf, gausspf, fdpf, dcpf, hepf, cnr, radial, pf3p")
        print(f"  Economics   : dcopf, acopf, sdpopf, uopf, scopf, opf3p, mpopf, stopf, market, lmp")
        print(f"  Grid Tools  : contingency, se, cpf, audit, varplan")
        print("\nExamples:")
        print("  acpf case14 1e-5 excel")
        print("  hepf case9 docx")
        print("  hepf --help docx")
        print("="*80 + "\n")
        return
        
    analysis = sys.argv[1].lower()
    run_cli_command(analysis, sys.argv[2:])

if __name__ == "__main__":
    main()
